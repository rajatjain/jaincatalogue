#!/usr/bin/env python

import argparse
import io
import logging
import os
import shutil
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import Pt
import docx2pdf
import fitz  # PyMuPDF
from PIL import Image
from tqdm import tqdm

from google.cloud import vision

"""
The tool accepts a PDF file and converts it into a fully translated
.docx file. Here are the steps:

  - PyMuPDF (fitz) is used to extract pages as images directly from PDF (no temp files)
  - Google Vision APIs are used to extract text from each page image
  - All extracted text is combined into a final Microsoft Word (docx) file

ADVANTAGES of this approach:
  - No intermediate image files stored on disk
  - Much faster PDF processing with PyMuPDF
  - No external dependencies (imagemagick, ghostscript)
  - Direct memory-based image processing

NOTE:
This script uses Google Vision APIs which is chargeable by Google.
Please be aware of this before using this script. Check the prices here
to estimate your cost:
https://cloud.google.com/vision/pricing#prices
"""

"""
Requirements:
  - Google Cloud Account
  - Python 3.11 or above

Setup:
  - Create a Google Cloud Account
  - Create a new project from https://console.cloud.google.com
  - Setup credentials and download the credentials JSON file
    https://cloud.google.com/docs/authentication/getting-started
  - Enable Vision API: https://cloud.google.com/vision/docs/setup
  - Install Google Cloud SDK and login:
    gcloud auth application-default login

  - Install required Python libraries:
    pip install PyMuPDF pillow python-docx google-cloud-vision docx2pdf tqdm

"""

# Default values - will be overridden by command line args
DEFAULT_BASE_FOLDER = "%s/Documents/A/ocr" % os.getenv("HOME")

# Global variables set by parse_args()
FNAME_PREFIX = None
BASE_FILE = None
OUTPUT_FOLDER = None
START_PAGE = None
END_PAGE = None
TOP_CROP = None
BOTTOM_CROP = None

vision_client = vision.ImageAnnotatorClient()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(filename)s:%(lineno)d - %(message)s'
)
logger = logging.getLogger(__name__)

def parse_args():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='Extract text from PDF using Google Vision API'
    )
    parser.add_argument(
        '--filename', '-f',
        type=str,
        required=True,
        help='Full path to PDF file (including .pdf extension)'
    )
    parser.add_argument(
        '--start-page', '-s',
        type=int,
        help='Starting page number (1-based)'
    )
    parser.add_argument(
        '--end-page', '-e',
        type=int,
        help='Ending page number (1-based)'
    )
    parser.add_argument(
        '--top-crop', '-t',
        type=float,
        default=0.0,
        help='Percentage of page height to crop from top (0.0-100.0)'
    )
    parser.add_argument(
        '--bottom-crop', '-b',
        type=float,
        default=0.0,
        help='Percentage of page height to crop from bottom (0.0-100.0)'
    )
    
    args = parser.parse_args()
    
    # Set global variables
    global FNAME_PREFIX, BASE_FILE, OUTPUT_FOLDER, START_PAGE, END_PAGE, TOP_CROP, BOTTOM_CROP
    
    BASE_FILE = args.filename
    # Extract filename without extension for prefix
    base_name = os.path.splitext(os.path.basename(BASE_FILE))[0]
    FNAME_PREFIX = base_name
    
    # Set output folder in same directory as input file
    file_dir = os.path.dirname(BASE_FILE)
    OUTPUT_FOLDER = os.path.join(file_dir, f"output_{FNAME_PREFIX}")
    
    START_PAGE = args.start_page - 1 if args.start_page else None  # Convert to 0-based
    END_PAGE = args.end_page - 1 if args.end_page else None  # Convert to 0-based
    TOP_CROP = args.top_crop
    BOTTOM_CROP = args.bottom_crop
    
    return args

def init():
    """Create clean output folder"""
    # Remove existing output folder if it exists
    if os.path.exists(OUTPUT_FOLDER):
        shutil.rmtree(OUTPUT_FOLDER)
        logger.info(f"Removed existing output folder: {OUTPUT_FOLDER}")
    
    # Create fresh output folder
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    logger.info(f"Created output folder: {OUTPUT_FOLDER}")

def extract_page_as_image(pdf_document, page_num):
    """Extract a single PDF page as PIL Image object with optional cropping"""
    try:
        page = pdf_document[page_num]
        # Render page as image with high resolution (300 DPI equivalent)
        mat = fitz.Matrix(2.0, 2.0)  # 2.0 = 144 DPI, adjust for quality vs speed
        pix = page.get_pixmap(matrix=mat)

        # Convert to PIL Image
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        
        # Apply cropping if specified (percentage-based)
        if TOP_CROP > 0 or BOTTOM_CROP > 0:
            width, height = img.size
            
            # Convert percentages to pixels
            top_pixels = int((TOP_CROP / 100.0) * height)
            bottom_pixels = int((BOTTOM_CROP / 100.0) * height)
            
            left = 0
            top = top_pixels
            right = width
            bottom = height - bottom_pixels
            
            # Ensure crop bounds are valid
            if bottom > top and right > left:
                img = img.crop((left, top, right, bottom))
                logger.debug(f"Cropped page {page_num + 1}: top={TOP_CROP}% ({top_pixels}px), bottom={BOTTOM_CROP}% ({bottom_pixels}px)")
            else:
                logger.warning(f"Invalid crop bounds for page {page_num + 1}, skipping crop")
        
        return img
    except Exception as e:
        logger.error(f"Error extracting page {page_num + 1}: {e}")
        return None

def detect_text_from_image(img, page_num, progress_bar):
    """Extract text from PIL Image using Google Vision API"""
    try:
        # Convert PIL Image to bytes
        img_byte_array = io.BytesIO()
        img.save(img_byte_array, format='PNG')
        img_bytes = img_byte_array.getvalue()

        # Create Vision API image object
        image = vision.Image(content=img_bytes)

        # Retry mechanism for Vision API
        success = False
        text_detection_response = None

        for attempt in range(2):
            try:
                text_detection_response = vision_client.document_text_detection(image=image)
                success = True
                break
            except Exception as e:
                if attempt == 0:
                    time.sleep(1)  # Brief pause before retry
                else:
                    logger.error(f"Failed to process page {page_num + 1} after retries: {e}")

        if success and text_detection_response:
            document = text_detection_response.full_text_annotation
            if document:
                paragraphs = []
                if document.pages:
                    page = document.pages[0]
                    for block in page.blocks:
                        for paragraph in block.paragraphs:
                            paragraph_text = ""
                            for word in paragraph.words:
                                word_text = ""
                                for symbol in word.symbols:
                                    word_text += symbol.text
                                paragraph_text += word_text + " "
                            if paragraph_text.strip():
                                paragraphs.append(paragraph_text.strip())
                text = "\n\n".join(paragraphs)
            else:
                text = ''
        else:
            text = ''

        progress_bar.update(1)
        return text, page_num

    except Exception as e:
        logger.error(f"Error processing page {page_num + 1}: {e}")
        progress_bar.update(1)
        return '', page_num

def process_single_page(pdf_path, page_num, progress_bar):
    """Process a single page: extract image and get text"""
    try:
        # Open PDF document for this thread (fitz is thread-safe for read operations)
        pdf_document = fitz.open(pdf_path)

        # Extract page as image
        img = extract_page_as_image(pdf_document, page_num)
        pdf_document.close()

        if img is None:
            progress_bar.update(1)
            return '', page_num

        # Extract text from image
        text, _ = detect_text_from_image(img, page_num, progress_bar)
        return text, page_num

    except Exception as e:
        logger.error(f"Error processing page {page_num + 1}: {e}")
        progress_bar.update(1)
        return '', page_num

def extract_text_from_pdf_parallel():
    """Extract text from all PDF pages in parallel"""

    logger.info(f"Opening PDF: {BASE_FILE}")

    try:
        # Open PDF to get page count
        pdf_document = fitz.open(BASE_FILE)
        total_pages = len(pdf_document)
        pdf_document.close()

        # Determine page range
        start_page = START_PAGE if START_PAGE is not None else 0
        end_page = END_PAGE if END_PAGE is not None else total_pages - 1
        
        # Validate page range
        start_page = max(0, min(start_page, total_pages - 1))
        end_page = max(start_page, min(end_page, total_pages - 1))
        
        pages_to_process = list(range(start_page, end_page + 1))
        logger.info(f"Processing pages {start_page + 1}-{end_page + 1} ({len(pages_to_process)} pages out of {total_pages} total)...")

        # Process pages in parallel
        max_workers = min(8, len(pages_to_process))  # Balance between speed and API limits
        page_texts = {}

        # Create progress bar
        with tqdm(total=len(pages_to_process), desc="Extracting text", unit="page") as progress_bar:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all page processing tasks
                futures = {
                    executor.submit(process_single_page, BASE_FILE, page_num, progress_bar): page_num
                    for page_num in pages_to_process
                }

                # Collect results
                for future in as_completed(futures):
                    text, page_num = future.result()
                    page_texts[page_num] = text

        logger.info(f"Text extraction completed for {len(pages_to_process)} pages!")

        # Return texts in correct page order
        ordered_texts = [page_texts[i] for i in pages_to_process]
        return ordered_texts

    except Exception as e:
        logger.error(f"Error opening PDF file: {e}")
        return []

def create_docx_from_texts(page_texts):
    """Create DOCX document from extracted texts"""
    if not page_texts:
        logger.warning("No text extracted to create document!")
        return None

    base_fname = FNAME_PREFIX
    final_fname = os.path.join(OUTPUT_FOLDER, f"{base_fname}_extracted.docx")

    logger.info(f"Creating DOCX with {len(page_texts)} pages...")

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'NotoSansDevanagari-Regular'
    style.font.size = Pt(10)

    # Use tqdm for DOCX creation progress
    with tqdm(total=len(page_texts), desc="Creating DOCX", unit="page") as progress_bar:
        for i, text in enumerate(page_texts, 1):
            if text.strip():  # Only add non-empty pages
                # Add page number header
                document.add_paragraph(f"--- Page {i} ---", style=style)
                document.add_paragraph(text, style=style)
            else:
                document.add_paragraph(f"--- Page {i} (No text detected) ---", style=style)

            document.add_page_break()
            progress_bar.update(1)

    document.save(final_fname)
    logger.info(f"DOCX file saved: {final_fname}")
    return final_fname

def save_text_files(page_texts):
    """Save extracted text to individual page files and a combined file"""
    if not page_texts:
        logger.warning("No text to save!")
        return None, None
    
    # Save individual page files
    page_files = []
    start_page_num = (START_PAGE if START_PAGE is not None else 0) + 1
    
    logger.info(f"Saving individual page text files...")
    for i, text in enumerate(page_texts):
        actual_page_num = start_page_num + i
        page_filename = os.path.join(OUTPUT_FOLDER, f"page_{actual_page_num:03d}.txt")
        
        try:
            with open(page_filename, 'w', encoding='utf-8') as f:
                f.write(text)
            page_files.append(page_filename)
            logger.debug(f"Saved page {actual_page_num}: {page_filename}")
        except Exception as e:
            logger.error(f"Error saving page {actual_page_num}: {e}")
    
    # Save combined file
    combined_filename = os.path.join(OUTPUT_FOLDER, f"{FNAME_PREFIX}_all_pages.txt")
    logger.info(f"Saving combined text file...")
    
    try:
        with open(combined_filename, 'w', encoding='utf-8') as f:
            for i, text in enumerate(page_texts):
                actual_page_num = start_page_num + i
                f.write(f"{'='*50}\n")
                f.write(f"Page {actual_page_num}\n")
                f.write(f"{'='*50}\n\n")
                f.write(text)
                f.write("\n\n")
        
        logger.info(f"Combined file saved: {combined_filename}")
        return page_files, combined_filename
        
    except Exception as e:
        logger.error(f"Error saving combined file: {e}")
        return page_files, None

def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF"""
    if not docx_path:
        return None

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(OUTPUT_FOLDER, f"{base_name}.pdf")

    logger.info("Converting DOCX to PDF...")
    try:
        docx2pdf.convert(docx_path, pdf_path)
        logger.info(f"PDF file saved: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"Error converting to PDF: {e}")
        return None

def main():
    """Main execution function"""
    start = time.time()
    
    # Parse command line arguments first
    parse_args()

    logger.info("=== PDF Text Extraction Tool (PyMuPDF + Vision API) ===")
    logger.info(f"Processing file: {BASE_FILE}")
    if START_PAGE is not None or END_PAGE is not None:
        start_display = (START_PAGE + 1) if START_PAGE is not None else "first"
        end_display = (END_PAGE + 1) if END_PAGE is not None else "last"
        logger.info(f"Page range: {start_display} to {end_display}")

    # Check if input file exists
    if not os.path.exists(BASE_FILE):
        logger.error(f"Input file not found: {BASE_FILE}")
        return

    # Initialize
    logger.info("Initializing...")
    init()

    try:
        # Step 1: Extract text from all pages
        logger.info("--- Step 1: Extracting text from PDF pages ---")
        page_texts = extract_text_from_pdf_parallel()

        if not page_texts:
            logger.warning("No text extracted. Exiting...")
            return

        # Step 2: Save extracted text to files
        logger.info("--- Step 2: Saving extracted text to files ---")
        page_files, combined_file = save_text_files(page_texts)

        # Summary
        end = time.time()
        total_time = end - start
        minutes = int(total_time // 60)
        seconds = int(total_time % 60)

        logger.info("=== PROCESSING COMPLETED ===")
        logger.info(f"Total processing time: {minutes}m {seconds}s ({total_time:.1f} seconds)")
        logger.info(f"Output files created in: {OUTPUT_FOLDER}")
        
        if page_files:
            logger.info(f"  - Individual page files: {len(page_files)} files (page_XXX.txt)")
        if combined_file:
            logger.info(f"  - Combined file: {os.path.basename(combined_file)}")

        # Calculate some stats
        total_chars = sum(len(text) for text in page_texts)
        avg_time_per_page = total_time / len(page_texts) if page_texts else 0

        logger.info("Statistics:")
        logger.info(f"  - Pages processed: {len(page_texts)}")
        logger.info(f"  - Total characters extracted: {total_chars:,}")
        logger.info(f"  - Average time per page: {avg_time_per_page:.1f} seconds")

    except KeyboardInterrupt:
        logger.info("Process interrupted by user. Cleaning up...")
    except Exception as e:
        logger.error(f"Unexpected error: {e}")

if __name__ == '__main__':
    main()